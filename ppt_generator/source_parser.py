"""Generic Excel and Word text extraction for the desktop import workspace."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


MAX_PREVIEW_CHARS = 200_000


@dataclass(frozen=True)
class ParsedSource:
    path: Path
    kind: str
    content: str
    section_count: int
    truncated: bool = False


def _limit_preview(content: str) -> tuple[str, bool]:
    if len(content) <= MAX_PREVIEW_CHARS:
        return content, False
    return content[:MAX_PREVIEW_CHARS] + "\n\n……内容过长，预览已截断……", True


def _parse_excel(path: Path) -> ParsedSource:
    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            lines.append(f"【工作表：{worksheet.title}】")
            for row in worksheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    lines.append("\t".join(cells))
            lines.append("")
        content, truncated = _limit_preview("\n".join(lines).strip())
        return ParsedSource(path, "Excel", content, len(workbook.worksheets), truncated)
    finally:
        workbook.close()


def _parse_word(path: Path) -> ParsedSource:
    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"【表格 {table_index}】")
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    content, truncated = _limit_preview("\n".join(lines).strip())
    return ParsedSource(path, "Word", content, len(document.paragraphs) + len(document.tables), truncated)


def parse_source(path: str | Path) -> ParsedSource:
    source = Path(path).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(f"资料文件不存在：{source}")
    suffix = source.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        return _parse_excel(source)
    if suffix == ".docx":
        return _parse_word(source)
    raise ValueError("当前只支持 .xlsx、.xlsm 和 .docx 资料")
