from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from ppt_generator import (
    AssetRecord,
    ExcelMappingRule,
    NavigationItem,
    PptProject,
    PresentationStyle,
    SourceRecord,
    load_project,
    parse_source,
    save_project,
)


class ProjectAndSourceTests(unittest.TestCase):
    def test_project_round_trip(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path(__file__).parent) as temp_dir:
            project = PptProject(
                project_name="筒形壳体方案",
                template_path="template.pptx",
                manifest_path="manifest.json",
                values={"title": "测试", "table": [["A", "B"]]},
                enabled_modules=["cover"],
                module_order=["cover"],
                sources=[SourceRecord("input.xlsx", "Excel", "已解析")],
                assets=[AssetRecord("product.png", "产品图", "product_image")],
                excel_path="input.xlsx",
                excel_mappings=[ExcelMappingRule("参数", "B2", "project_title")],
                presentation_style=PresentationStyle(
                    navigation_height=0.60,
                    navigation_background="#DDEEFF",
                    navigation_font_size=13.0,
                    navigation_items=[
                        NavigationItem("设备设计", ["equipment_overview"]),
                        NavigationItem("检测效果", ["inspection_result"]),
                    ],
                ),
            )
            project.no_cad_scene = {
                "schemaVersion": "no-cad-equipment-scene/v2",
                "projectName": "候选图持久化测试",
            }
            project.ai_image_batches = [
                {
                    "schemaVersion": "project-ai-image-batch/v1",
                    "projectId": project.project_id,
                    "generationTarget": {
                        "targetId": "M04",
                        "targetHash": "a" * 64,
                    },
                    "batch": {"batchId": "batch-01", "candidates": []},
                }
            ]
            path = Path(temp_dir) / "sample.kyppt.json"

            save_project(project, path)
            restored = load_project(path)

            self.assertEqual(restored.project_name, project.project_name)
            self.assertEqual(restored.project_id, project.project_id)
            self.assertEqual(restored.no_cad_scene, project.no_cad_scene)
            self.assertEqual(restored.ai_image_batches, project.ai_image_batches)
            self.assertEqual(restored.values, project.values)
            self.assertEqual(restored.sources[0].content, "已解析")
            self.assertEqual(restored.assets[0].slot_key, "product_image")
            self.assertEqual(restored.excel_path, "input.xlsx")
            self.assertEqual(restored.excel_mappings[0].source_range, "B2")
            self.assertEqual(restored.presentation_style.navigation_height, 0.60)
            self.assertEqual(
                restored.presentation_style.navigation_background,
                "#DDEEFF",
            )
            self.assertEqual(restored.presentation_style.navigation_font_size, 13.0)
            self.assertEqual(
                [item.name for item in restored.presentation_style.navigation_items],
                ["设备设计", "检测效果"],
            )
            self.assertEqual(
                restored.presentation_style.navigation_items[1].module_keys,
                ["inspection_result"],
            )
            raw = json.loads(path.read_text(encoding="utf-8"))
            self.assertEqual(raw["schema_version"], 5)
            self.assertFalse(path.with_suffix(path.suffix + ".tmp").exists())

    def test_legacy_project_gets_stable_id_when_saved_as_schema_v5(self) -> None:
        restored = PptProject.from_dict(
            {
                "schema_version": 4,
                "project_name": "legacy",
                "values": {},
            }
        )
        self.assertRegex(restored.project_id, r"^[0-9a-f]{32}$")
        saved = restored.to_dict()
        self.assertEqual(saved["schema_version"], 5)
        self.assertEqual(saved["project_id"], restored.project_id)

    def test_rejects_candidate_batch_owned_by_another_project(self) -> None:
        project = PptProject()
        raw = project.to_dict()
        raw["ai_image_batches"] = [
            {
                "projectId": "f" * 32,
                "generationTarget": {
                    "targetId": "overview",
                    "targetHash": "a" * 64,
                },
                "batch": {"batchId": "foreign", "candidates": []},
            }
        ]
        with self.assertRaises(ValueError):
            PptProject.from_dict(raw)

    def test_legacy_style_loads_new_default_navigation(self) -> None:
        style = PresentationStyle.from_dict(
            {
                "navigation_height": 0.52,
                "navigation_background": "#FFFFFF",
            }
        )
        self.assertEqual(
            [item.name for item in style.navigation_items],
            ["公司简介", "工艺分析", "设备设计", "检测效果", "系统介绍"],
        )
        assigned = {
            module_key
            for item in style.navigation_items
            for module_key in item.module_keys
        }
        self.assertEqual(len(assigned), 16)
        self.assertEqual(style.navigation_index_for("inspection_result"), 3)
        self.assertIsNone(style.navigation_font_size)
        self.assertEqual(style.resolved_navigation_font_size(), 10.0)
        style.navigation_height = 0.72
        self.assertEqual(style.resolved_navigation_font_size(), 14.0)
        style.navigation_font_size = 12.0
        self.assertEqual(style.resolved_navigation_font_size(), 12.0)

    def test_rejects_invalid_presentation_style(self) -> None:
        project = PptProject()
        project.presentation_style.navigation_background = "red"
        with self.assertRaisesRegex(ValueError, "#RRGGBB"):
            project.to_dict()

        project.presentation_style = PresentationStyle(
            navigation_items=[
                NavigationItem("设备设计", ["equipment_overview"]),
                NavigationItem("设备设计", ["inspection_result"]),
            ]
        )
        with self.assertRaisesRegex(ValueError, "名称不能重复"):
            project.to_dict()

        project.presentation_style = PresentationStyle(
            navigation_items=[
                NavigationItem("设备设计", ["equipment_overview"]),
                NavigationItem("检测效果", ["equipment_overview"]),
            ]
        )
        with self.assertRaisesRegex(ValueError, "归属多个"):
            project.to_dict()

        with self.assertRaisesRegex(ValueError, "1～7"):
            PresentationStyle(
                navigation_items=[NavigationItem(f"栏目{index}") for index in range(8)]
            ).validate()

        with self.assertRaisesRegex(ValueError, "10 个字符"):
            PresentationStyle(
                navigation_items=[NavigationItem("这是一个超过十个字符的导航栏目")]
            ).validate()

        with self.assertRaisesRegex(ValueError, "9～16"):
            PresentationStyle(navigation_font_size=17.0).validate()

    def test_parses_excel_and_word(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path(__file__).parent) as temp_dir:
            directory = Path(temp_dir)
            excel_path = directory / "input.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "设备参数"
            sheet.append(["参数", "值"])
            sheet.append(["节拍", "30 pcs/min"])
            workbook.save(excel_path)

            word_path = directory / "input.docx"
            document = Document()
            document.add_heading("技术要求", level=1)
            document.add_paragraph("检测筒形壳体外观缺陷")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "精度"
            table.cell(0, 1).text = "0.8 mm"
            document.save(word_path)

            excel = parse_source(excel_path)
            word = parse_source(word_path)

            self.assertEqual(excel.kind, "Excel")
            self.assertIn("设备参数", excel.content)
            self.assertIn("30 pcs/min", excel.content)
            self.assertEqual(word.kind, "Word")
            self.assertIn("检测筒形壳体外观缺陷", word.content)
            self.assertIn("0.8 mm", word.content)

    def test_rejects_unsupported_source(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path(__file__).parent) as temp_dir:
            path = Path(temp_dir) / "input.txt"
            path.write_text("text", encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "只支持"):
                parse_source(path)


if __name__ == "__main__":
    unittest.main()
